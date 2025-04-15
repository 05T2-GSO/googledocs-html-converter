import os
import re
import requests
from pathlib import Path
from docx import Document
from docx.text.run import Run
from urllib.parse import urlparse
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATE_FILE = "template.html"

# Extracts the title from the html prior to downloading the document wwith python-docx
def get_google_doc_title_from_html(link):
    response = requests.get(link)
    response.raise_for_status()
    match = re.search(r"<title>(.*?)</title>", response.text)
    if match:
        title = match.group(1).replace(" - Google Docs", "").strip()
        return title or "Untitled_Document"
    return "Untitled_Document"


# This function is used to get the alignment style of a paragraph
def get_alignment_style(para):
    alignment = para.alignment
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return ' style="text-align:center;"'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return ' style="text-align:right;"'
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return ' style="text-align:justify;"'
    return ''

# This function is used to get the export link for a Google Docs document
def get_export_link(share_link):
    
    match = re.search(r"/document/d/([^/]+)", share_link)
    if not match:   
        raise ValueError("Invalid Google Docs link.")
    doc_id = match.group(1)
    return f"https://docs.google.com/document/d/{doc_id}/export?format=docx"

# This function is used to download the Google Docs document in .docx format
def download_docx(export_url, save_path):
    response = requests.get(export_url)
    response.raise_for_status()
    with open(save_path, 'wb') as f:
        f.write(response.content)

# This function is used to style the text in the document
# based on the formatting of the runs
def style_run(run: Run):
    text = run.text
    if not text:
        return ""

    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text

# This function is used to convert the font size to a corresponding HTML tag
def pt_to_tag(pt_value):
    if pt_value >= 32:
        return "h1"
    elif pt_value >= 24:
        return "h2"
    elif pt_value >= 18:
        return "h3"
    elif pt_value >= 14:
        return "h4"
    else:
        return "p"

# This function is used to convert the .docx document to HTML
def convert_docx_to_html(docx_path):
    doc = Document(docx_path)
    html_parts = []

    in_list = False
    list_tag = ""

    for para in doc.paragraphs:
        style = para.style.name.lower()
        styled_text = ''.join(style_run(run) for run in para.runs).strip()

        if not styled_text:
            continue

        align_attr = get_alignment_style(para)

        # Detect heading styles
        if "heading 1" in style:
            html_parts.append(f"<h1{align_attr}>{styled_text}</h1>")
        elif "heading 2" in style:
            html_parts.append(f"<h2{align_attr}>{styled_text}</h2>")
        elif "heading 3" in style:
            html_parts.append(f"<h3{align_attr}>{styled_text}</h3>")
        elif "list" in style or para.style.name.startswith("List"):
            if not in_list:
                list_tag = "ol" if "number" in style else "ul"
                html_parts.append(f"<{list_tag}>")
                in_list = True
            html_parts.append(f"<li>{styled_text}</li>")
        else:
            # If previously in a list, close it
            if in_list:
                html_parts.append(f"</{list_tag}>")
                in_list = False

            # Use font size to guess tag if not a heading
            max_font_pt = 0
            for run in para.runs:
                if run.font.size:
                    pt = run.font.size.pt
                    if pt > max_font_pt:
                        max_font_pt = pt

            tag = pt_to_tag(max_font_pt) if max_font_pt else "p"
            html_parts.append(f"<{tag}{align_attr}>{styled_text}</{tag}>")

    # If list still open
    if in_list:
        html_parts.append(f"</{list_tag}>")

    return "\n".join(html_parts)

# This function is used to apply the HTML template to the converted content
# and return the final HTML
def apply_template(html_body, template_path, doc_title="Document"):
    with open(template_path, 'r', encoding="utf-8") as f:
        template = f.read()
    return template.replace("{{ content }}", html_body).replace("{{ title }}", doc_title)

# This function is used to get the title of the document
# from the document properties
def get_doc_title(docx_path):
    doc = Document(docx_path)
    return doc.core_properties.title or "Untitled_Document"

def save_output(html, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    html_path = os.path.join(output_dir, "index.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

def process_google_doc(link):
    export_url = get_export_link(link)
    tmp_docx = "temp.docx"
    download_docx(export_url, tmp_docx)

    doc_title = get_google_doc_title_from_html(link)    
    output_dir = Path("output") / doc_title.replace(" ", "_")
    output_dir.mkdir(parents=True, exist_ok=True)

    html_body = convert_docx_to_html(tmp_docx)
    full_html = apply_template(html_body, TEMPLATE_FILE, doc_title)
    save_output(full_html, output_dir)

    os.remove(tmp_docx)
    print(f"Completed: Document compiled to: {output_dir}/index.html")

# Example run
if __name__ == "__main__":
    example_link = input("Paste your Google Docs share link: ").strip()
    process_google_doc(example_link)
